from django.http import HttpResponse
from django_xhtml2pdf.utils import generate_pdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches
from io import StringIO
from lxml import etree

docx = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

# https://python-docx.readthedocs.io/en/latest/#what-it-can-do

def render_doc(context, **kwargs):
    """
    """

    # # https://stackoverflow.com/a/24122313/185820
    # document = Document()
    # # Head
    # task = ''
    # contract = context['item']
    # if contract.task:
    #     task = contract.task
    # title = document.add_heading(
    #     'ACLARK.NET, LLC %s AGREEMENT PREPARED FOR:' % task, level=1)
    # title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # if contract.client:
    #     client_name = document.add_heading(contract.client.name, level=1)
    #     client_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     client_address = document.add_heading(contract.client.address, level=1)
    #     client_address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # parser = etree.HTMLParser()  # http://lxml.de/parsing.html
    # tree = etree.parse(StringIO(contract.body), parser)
    # # Body
    # for element in tree.iter():
    #     if element.tag == 'h2':
    #         document.add_heading(element.text, level=2)
    #     elif element.tag == 'p':
    #         document.add_paragraph(element.text)
    # response = HttpResponse(content_type=DOC)
    # response['Content-Disposition'] = 'attachment; filename=download.docx'
    # document.save(response)
    # return response

    document = Document()

    filename = kwargs.get('filename')

    # logo = kwargs.get('logo')
    # document.add_picture(logo, height=Inches(0.33))

    item = context['item']
    document.add_heading(item.title, 0)

    p = document.add_paragraph(item.note)

    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    # document.add_paragraph(
    #     'first item in ordered list', style='List Number'
    # )

    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )

    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc
    # document.add_page_break()

    response = HttpResponse(content_type=docx)
    response['Content-Disposition'] = 'attachment; filename=%s' % filename
    document.save(response)
    return response


def render_pdf(context, **kwargs):
    filename = kwargs.get('filename')
    template = kwargs.get('template')
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'filename=%s' % filename
    return generate_pdf(template, context=context, file_object=response)
